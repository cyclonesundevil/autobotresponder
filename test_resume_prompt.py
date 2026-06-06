import unittest
from tempfile import TemporaryDirectory

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn

from resume_processor import add_hyperlink_run, build_resume_prompt, render_tailored_resume_on_template


class ResumePromptTests(unittest.TestCase):
    def test_hyperlink_run_creates_clickable_word_relationship(self):
        doc = Document()
        paragraph = doc.add_paragraph()

        add_hyperlink_run(paragraph, "GitHub", "https://github.com/cyclonesundevil")

        self.assertEqual(paragraph.text, "GitHub")
        self.assertTrue(any(rel.reltype == RT.HYPERLINK for rel in paragraph.part.rels.values()))
        self.assertEqual(paragraph._p.xpath(".//w:hyperlink")[-1].get(qn('w:history')), '1')

    def test_prompt_stays_factual_and_does_not_overstate(self):
        prompt = build_resume_prompt(
            email_body="Recruiter wants a Python/ML engineer",
            base_text="Experience: Python developer at Example Corp",
        )

        self.assertIn("Do not invent", prompt)
        self.assertIn("preserve the existing facts", prompt)
        self.assertIn("Do not invent schools, degrees, institutions, certifications, or names", prompt)
        self.assertNotIn("aggressively emphasize", prompt)
        self.assertNotIn("AI-focused engineer", prompt)

    def test_render_tailored_resume_keeps_template_paragraph_styles(self):
        with TemporaryDirectory() as tmpdir:
            template_path = f"{tmpdir}/template.docx"
            output_path = f"{tmpdir}/output.docx"

            template = Document()
            template.add_heading("Original Name", level=1)
            template.add_paragraph("Original summary")
            template.add_paragraph("Original bullet", style="List Paragraph")
            template.save(template_path)

            render_tailored_resume_on_template(
                template_path,
                "# Jose C. Ramirez\nUpdated summary\n- Updated bullet",
                output_path,
            )

            output = Document(output_path)
            self.assertEqual([p.text for p in output.paragraphs], [
                "Jose C. Ramirez",
                "Updated summary",
                "Updated bullet",
            ])
            self.assertEqual([p.style.name for p in output.paragraphs], [
                "Heading 1",
                "Normal",
                "List Paragraph",
            ])


if __name__ == "__main__":
    unittest.main()
