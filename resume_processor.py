import os
import re
from dotenv import load_dotenv
from google import genai
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

load_dotenv()
API_KEY = os.getenv("GEMINI_API_KEY")

if API_KEY:
    client = genai.Client(api_key=API_KEY)
else:
    client = None
    
# gemini-2.5-flash-lite: Google's smallest/cheapest model, built for at-scale usage
MODEL_ID = 'gemini-2.5-flash-lite'


def build_resume_prompt(email_body, base_text):
    """Build a conservative prompt that tailors the resume without inventing facts."""
    return f"""
    You are a careful resume editor.
    Tailor the resume to the recruiter email using only the facts already present in the base resume.

    CRITICAL RULES:
    - preserve the existing facts, chronology, job titles, dates, employers, and achievements from the base resume.
    - preserve the existing facts and wording style of the base resume wherever possible.
    - Do not invent schools, degrees, institutions, certifications, or names.
    - Do not invent new companies, roles, dates, technologies, metrics, or accomplishments.
    - Do not add unsupported claims or make the resume sound more senior than the source material.
    - If a section, institution, school, degree, or detail is not present in the base resume, omit it instead of creating it.
    - If the recruiter email mentions skills or keywords, highlight them only when they are already supported by the existing resume content.
    - Keep the tone professional and factual; avoid exaggerated AI-only framing or suspicious overstatement.
    - Output the tailored resume in PURE MARKDOWN format.

    Recruiter Email:
    {email_body}

    Base Resume:
    {base_text}
    """


def extract_text_from_docx(file_path):
    """Reads paragraphs and tables from a word doc to ensure no content is missed."""
    doc = Document(file_path)
    full_text = []
    # Extract from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    # Extract from tables (resumes often use tables for layout)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text)
    return '\n'.join(full_text)

def is_recruiter_opportunity(email_body):
    """
    Uses Gemini to determine if the email is a legitimate recruiter job opportunity.
    Returns True if it is, False otherwise.
    """
    if not API_KEY:
        return True # Default to True if AI is not available to avoid missing any

    prompt = f"""
    Analyze the following email and determine if it is a legitimate job opportunity, 
    a recruiter reaching out about a role, or a forwarded recruitment email.
    
    CRITERIA FOR 'TRUE':
    - Recruiters (internal or agency) reaching out about current or future roles.
    - Automated alerts from job boards (LinkedIn, Indeed) about specific matching jobs.
    - Forwarded emails that contain a job description or recruiter reach-out.
    - Mention of specific titles (Engineer, Developer, etc.), companies, or interview requests.

    CRITERIA FOR 'FALSE':
    - Generic newsletters, marketing spam, or promotional material (e.g., Domino's, SoFi, Total Wine).
    - Personal non-work correspondence.
    - "Work from home" scams or generic "earn money fast" emails.

    Respond with ONLY the word 'TRUE' if it is a legitimate recruiting/job opportunity, 
    or 'FALSE' if it is junk/irrelevant.
    
    Email:
    {email_body}
    """
    try:
        response = client.models.generate_content(model=MODEL_ID, contents=prompt).text.strip()
        return "TRUE" in response.upper()
    except Exception as e:
        print(f"Error classifying email: {e}")
        return True # Default to True on error

def extract_forward_to_email(email_body, sender_email):
    """
    Uses Gemini to determine if there is an explicit email address to forward
    the resume/application to, instead of the original sender.
    Returns the target email address, or the original sender_email if none found.
    """
    if not API_KEY:
        return sender_email
        
    prompt = f"""
    Analyze the following email and determine if the sender is asking for a resume or application 
    to be sent to a SPECIFIC email address that is NOT the sender's address and NOT my own email addresses (cyclsun@gmail.com or cyclonesundevil@gmail.com).

    CRITICAL RULES:
    1. DO NOT return my own email addresses: cyclsun@gmail.com or cyclonesundevil@gmail.com.
    2. DO NOT return LinkedIn email addresses (e.g., matching-reply@linkedin.com) or LinkedIn profile URLs.
    3. IF THIS IS A FORWARDED EMAIL (starts with "Forwarded message" or similar): 
       - Look for the "From:" line in the forwarded header to find the ORIGINAL recruiter's email.
       - Use that original email as the target unless there is an even more specific instruction to "forward" to another address.
    4. ONLY return a specific email address if there is an explicit instruction to "forward", "send", or "email" the resume to it, or if it's the original sender of a forwarded message.
    5. If the email just says "apply on LinkedIn" or provides a link to an application portal, respond with 'NONE'.

    Look for phrases like:
    - "Please forward your resume to hr@example.com"
    - "Send your details to recruiting@somecompany.com"
    - "Reply to hiringmanager@domain.com"
    - "If interested, email your resume to jobs@test.com"

    If you find a specific, valid target email address (not my own and not generic LinkedIn),
    respond with ONLY that email address (e.g., hr@example.com).
    
    If there is no specific target address found, respond with 'NONE'.

    Email:
    {email_body}
    """
    try:
        response = client.models.generate_content(model=MODEL_ID, contents=prompt).text.strip()
        # Clean up the output to make sure it's just the email or NONE
        response = response.strip("<> \n\r\t*")
        
        if response.upper() == 'NONE' or '@' not in response:
            return sender_email
            
        return response
    except Exception as e:
        print(f"Error extracting forward-to email: {e}")
        return sender_email

URL_RE = re.compile(r'(?P<url>https?://[^\s)<>]+|www\.[^\s)<>]+)')
MARKDOWN_LINK_RE = re.compile(r'\[([^\]]+?)\]\((https?://[^)]+)\)')


def normalize_url(url):
    """Convert bare website URLs into clickable http(s) links."""
    cleaned = url.strip().strip(').,;:')
    if cleaned.startswith('www.'):
        return f'https://{cleaned}'
    return cleaned


def add_hyperlink_run(paragraph, text, url):
    """Create a real Word hyperlink run for the given URL."""
    clean_url = normalize_url(url)
    rels = paragraph.part.rels
    r_id = rels._next_rId
    rels.add_relationship(RT.HYPERLINK, clean_url, r_id, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    hyperlink.set(qn('w:history'), '1')

    run = OxmlElement('w:r')
    run_props = OxmlElement('w:rPr')
    style = OxmlElement('w:rStyle')
    style.set(qn('w:val'), 'Hyperlink')
    run_props.append(style)
    run.append(run_props)

    text_element = OxmlElement('w:t')
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_text_with_links(paragraph, text):
    """Append plain text plus clickable URL segments to a paragraph."""
    cleaned = text.replace('**', '').replace('`', '')

    cursor = 0
    for match in MARKDOWN_LINK_RE.finditer(cleaned):
        if match.start() > cursor:
            paragraph.add_run(cleaned[cursor:match.start()])
        add_hyperlink_run(paragraph, match.group(1), match.group(2))
        cursor = match.end()

    tail = cleaned[cursor:]
    if not tail:
        return

    cursor = 0
    for match in URL_RE.finditer(tail):
        if match.start() > cursor:
            paragraph.add_run(tail[cursor:match.start()])
        display_text = normalize_url(match.group('url'))
        add_hyperlink_run(paragraph, display_text, match.group('url'))
        cursor = match.end()

    if cursor < len(tail):
        paragraph.add_run(tail[cursor:])


def ensure_hyperlink_style(doc):
    """Ensure generated DOCX files contain the Word Hyperlink style used by native links."""
    style_names = {style.name for style in doc.styles}
    if 'Hyperlink' not in style_names:
        style = doc.styles.add_style('Hyperlink', WD_STYLE_TYPE.CHARACTER)
        style.font.color.rgb = RGBColor(0, 102, 204)
        style.font.underline = True
    return doc


def create_template_document(base_resume_path):
    """Create a DOCX from the base resume template so its formatting survives."""
    doc = Document(base_resume_path)
    return ensure_hyperlink_style(doc)


def iter_resume_paragraphs(doc):
    """Return paragraphs that carry visible resume text, including table cells."""
    paragraphs = [paragraph for paragraph in doc.paragraphs if paragraph.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(paragraph for paragraph in cell.paragraphs if paragraph.text.strip())
    return paragraphs


def markdown_to_resume_lines(tailored_md):
    """Normalize Gemini markdown into one resume line per template paragraph."""
    lines = []
    for raw_line in tailored_md.replace("```markdown", "").replace("```", "").splitlines():
        line = raw_line.strip()
        if not line or line == "---":
            continue
        line = re.sub(r'^\s*#{1,6}\s+', '', line)
        line = re.sub(r'^\s*[-*]\s+', '', line)
        line = line.replace('**', '').replace('`', '').strip()
        if line:
            lines.append(line)
    return lines


def clear_paragraph_content(paragraph):
    """Remove runs/hyperlinks while keeping paragraph-level formatting intact."""
    for child in list(paragraph._p):
        if child.tag != qn('w:pPr'):
            paragraph._p.remove(child)


def replace_paragraph_text(paragraph, text):
    """Replace text inside an existing paragraph without changing its style/layout."""
    clear_paragraph_content(paragraph)
    add_text_with_links(paragraph, text)


def render_tailored_resume_on_template(base_resume_path, tailored_md, output_path):
    """Save tailored resume text into the existing base resume DOCX formatting."""
    doc = create_template_document(base_resume_path)
    template_paragraphs = iter_resume_paragraphs(doc)
    tailored_lines = markdown_to_resume_lines(tailored_md)

    if not tailored_lines:
        doc.save(output_path)
        return output_path

    for paragraph, line in zip(template_paragraphs, tailored_lines):
        replace_paragraph_text(paragraph, line)

    if len(tailored_lines) > len(template_paragraphs):
        fallback_style = template_paragraphs[-1].style if template_paragraphs else None
        for line in tailored_lines[len(template_paragraphs):]:
            paragraph = doc.add_paragraph(style=fallback_style)
            add_text_with_links(paragraph, line)

    doc.save(output_path)
    return output_path


def sanitize_tailored_markdown(tailored_md, base_text):
    """Drop invented education/school details that are not supported by the base resume."""
    base_text_lower = base_text.lower()
    sanitized_lines = []
    skip_education_section = False

    for line in tailored_md.splitlines():
        stripped = line.strip()

        if re.match(r'^(#{1,6}\s*)?education\b', stripped, re.I):
            skip_education_section = True
            continue

        if skip_education_section and (stripped.startswith('#') or stripped.startswith('##') or stripped.startswith('###') or stripped == '---'):
            skip_education_section = False

        if skip_education_section:
            continue

        if re.search(r'\b(bachelor|bachelor\'s|master|masteral|doctoral|associate|degree|university|college|school|stanford|berkeley)\b', stripped, re.I):
            if 'education' not in base_text_lower and 'degree' not in base_text_lower and 'university' not in base_text_lower and 'college' not in base_text_lower:
                continue

        sanitized_lines.append(line)

    return '\n'.join(sanitized_lines).strip()


def generate_tailored_resume_docx(email_body, base_resume_path, output_path):
    """
    1. Extracts job requirements from email
    2. Tailors base resume
    3. Saves tailored resume as a simple docx
    """
    if not API_KEY:
        print("GEMINI_API_KEY is not set. Cannot use AI.")
        return None

    print(f"Reading base resume from {base_resume_path}...")
    base_text = extract_text_from_docx(base_resume_path)

    print("Tailoring resume with Gemini (single call)...")
    combined_prompt = build_resume_prompt(email_body, base_text)
    tailored_md = client.models.generate_content(model=MODEL_ID, contents=combined_prompt).text
    tailored_md = sanitize_tailored_markdown(tailored_md, base_text)

    if re.search(r'\b(stanford|berkeley|bachelor|bachelor\'s|master|masteral|doctoral|associate|degree|university|college|school|30\+ years|30 years)\b', tailored_md, re.I):
        print("Suspicious hallucinated education/experience details detected; falling back to the verified base resume text.")
        tailored_md = base_text

    print(f"Generating formatted Word document at {output_path}...")
    render_tailored_resume_on_template(base_resume_path, tailored_md, output_path)
    print("Done generating resume!")
    return output_path

if __name__ == "__main__":
    # Test block
    print("This file contains the resume processing logic.")
