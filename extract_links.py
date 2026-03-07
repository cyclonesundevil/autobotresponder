from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

def get_links(path):
    document = Document(path)
    links = []
    
    # Check paragraphs
    for p in document.paragraphs:
        rels = p.part.rels
        for rel in rels.values():
            if rel.reltype == RT.HYPERLINK:
                links.append(f"{rel.target_ref}")
                
    # Check tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    rels = p.part.rels
                    for rel in rels.values():
                        if rel.reltype == RT.HYPERLINK:
                            links.append(f"{rel.target_ref}")
                            
    return list(set(links))

if __name__ == "__main__":
    resume_path = "C:\\Users\\cyclo\\OneDrive\\resumes\\cv.docx"
    print(f"Extracting links from {resume_path}...")
    try:
        links = get_links(resume_path)
        print("Found links:")
        for l in links:
            print(f"- {l}")
    except Exception as e:
        print(f"Error: {e}")
