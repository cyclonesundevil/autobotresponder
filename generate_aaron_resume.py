from resume_processor import generate_tailored_resume_docx, client, MODEL_ID, extract_text_from_docx
import os
from docx import Document
from dotenv import load_dotenv

load_dotenv()

def generate_custom_resume(email_body, base_resume_path, output_path):
    base_text = extract_text_from_docx(base_resume_path)
    
    # Hard-code the date correction to avoid AI hallucination
    # Original: Jan 2013- July 2014
    base_text = base_text.replace("Jan 2013 July 2014", "Jan 2013 - March 24, 2026")
    base_text = base_text.replace("Jan 2013- July 2014", "Jan 2013 - March 24, 2026")
    
    custom_prompt = f"""
    You are a careful resume editor.
    Tailor the resume to the recruiter email using only the facts already present in the base resume.

    CRITICAL RULES:
    - Preserve the existing employers, dates, titles, and achievements from the base resume.
    - Do not invent new companies, dates, technologies, or bullet points.
    - If the recruiter email mentions relevant keywords, emphasize them only through existing experience and skills already supported by the base resume.
    - Keep the resume factual, professional, and conservative.
    - Output the resume in PURE MARKDOWN format.

    Important personal details to preserve:
    - Name: Jose C. Ramirez
    - Location: Chandler, AZ
    - Phone: (480) 209-3709
    - Email: cyclsun@gmail.com
    - GitHub: https://github.com/cyclonesundevil

    Email from Recruiter:
    {email_body}

    Base Resume:
    {base_text}
    """
    
    print("Generating resume with custom prompt...")
    tailored_md = client.models.generate_content(model=MODEL_ID, contents=custom_prompt).text
    
    # Globally strip markdown code blocks if present
    tailored_md = tailored_md.replace("```markdown", "").replace("```", "").strip()
    
    # FORCE WEBSITE INCLUSIONS
    if "www.microcompit.com" not in tailored_md:
        tailored_md = tailored_md.replace("cyclsun@gmail.com", "cyclsun@gmail.com | www.microcompit.com")
    
    # STRIP ANY AI PARENTHETICAL NOTES
    import re
    tailored_md = re.sub(r'\(Note:.*?\)', '', tailored_md)
    # Ensure the date is exactly as requested
    tailored_md = tailored_md.replace("March 2026", "March 24, 2026")
    tailored_md = tailored_md.replace("March 24, 2026", "March 24, 2026") # Ensure it's not doubled

    print(f"Saving to {output_path}...")
    doc = Document()
    
    # MANUALLY ADD THE HEADER TO ENSURE IT IS ALWAYS THERE
    doc.add_heading("Jose C. Ramirez", level=1)
    doc.add_paragraph("Chandler, AZ | (480) 209-3709 | cyclsun@gmail.com | www.microcompit.com | https://github.com/cyclonesundevil")
    doc.add_paragraph("---")
    
    lines = tailored_md.split('\n')
    skip_header = True
    for line in lines:
        line = line.strip()
        if not line: continue
        
        # Skip the AI-generated header and conversational text
        if skip_header:
            if any(x in line for x in ["Jose C. Ramirez", "Chandler", "cyclsun", "github", "Here is your", "Markdown", "updated resume"]):
                continue
            if line == "---" or line == "Technical Skills":
                skip_header = False
                if line == "---": continue
        
        # Strip AI notes and enforce Wells Fargo dates
        import re
        line = re.sub(r'\(Note:.*?\)', '', line).strip()
        if "Wells Fargo" in line and ("Jan 2013" in line or "July 2014" in line):
            line = "Wells Fargo | Jan 2013 - March 24, 2026"

        if line.startswith('# '): doc.add_heading(line[2:], level=1)
        elif line.startswith('## '): doc.add_heading(line[3:], level=2)
        elif line.startswith('### '): doc.add_heading(line[4:], level=3)
        elif line.startswith('- '): doc.add_paragraph(line[2:].replace('**', ''), style='List Bullet')
        else: doc.add_paragraph(line.replace('**', ''))
            
    doc.save(output_path)
    print("Resume generated successfully!")
    return output_path

if __name__ == "__main__":
    with open('full_email_utf8.txt', 'r', encoding='utf-8') as f:
        email_body = f.read()
    generate_custom_resume(email_body, 'base_resume.docx', 'tailored_resume_final_v17.docx')
